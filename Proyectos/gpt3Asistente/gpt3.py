import os
import openai
import json
from docx import Document
from docx.shared import Cm

# Creación del documento
document = Document()

openai.api_key = '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'

response = openai.Completion.create(
  model="text-davinci-002",
  prompt="Realizar una instroduccion de los humanos",
  temperature=0.7,
  max_tokens=700,
  top_p=1,
  frequency_penalty=0,
  presence_penalty=0
)


print(response.choices[0].text.strip())
# Añadimos un párrafo
p = document.add_paragraph("Introducción")
p = document.add_paragraph(response.choices[0].text.strip())


document.save('ejemplo4.docx')